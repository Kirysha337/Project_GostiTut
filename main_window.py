import os
from datetime import date, datetime

from PyQt6.QtWidgets import (
    QMainWindow,
    QWidget,
    QHBoxLayout,
    QVBoxLayout,
    QFormLayout,
    QLabel,
    QPushButton,
    QStackedWidget,
    QLineEdit,
    QDialog,
    QTableWidget,
    QTableWidgetItem,
    QHeaderView,
    QMessageBox,
    QComboBox,
    QDateEdit,
    QSpinBox,
    QGroupBox,
    QTextEdit,
    QScrollArea,
    QFrame,
    QGridLayout,
    QInputDialog,
    QDoubleSpinBox,
    QCheckBox,
)
from PyQt6.QtGui import QPixmap
from PyQt6.QtCore import Qt, QDate, pyqtSignal

from config import (
    MAIN_IMAGE_PATH,
    SIDEBAR_COLOR,
    COLOR_FREE,
    COLOR_CLEANING,
    COLOR_OCCUPIED,
    COLOR_BOOKED,
    TITLE_FONT,
    SECTION_FONT,
    ROOM_FONT,
)
from crypto_utils import aes_encrypt, aes_decrypt
from db import db


class RoomTile(QLabel):
    """Плитка номера на главной странице."""

    clicked = pyqtSignal(object)

    def __init__(self, text, status, room_id, parent=None):
        super().__init__(text, parent)
        self.room_id = room_id
        self.status = status
        self.selected = False
        self.setFont(ROOM_FONT)
        self.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.setFixedHeight(48)
        self.apply_style()

    def mousePressEvent(self, event):
        self.clicked.emit(self)
        super().mousePressEvent(event)

    def set_status(self, status):
        self.status = status
        self.apply_style()

    def set_selected(self, selected: bool):
        self.selected = selected
        self.apply_style()

    def apply_style(self):
        color = COLOR_FREE
        if self.status == "уборка":
            color = COLOR_CLEANING
        elif self.status == "занят":
            color = COLOR_OCCUPIED
        elif self.status == "бронь":
            color = COLOR_BOOKED
        border = "2px solid #333" if self.selected else "none"
        self.setStyleSheet(
            f"background-color: {color}; border-radius: 8px; padding:6px; border:{border};"
        )


class MainWindow(QMainWindow):
    """Главное окно администратора."""

    def __init__(self, admin):
        super().__init__()
        self.admin = admin
        self.selected_tile = None

        self.setWindowTitle("ГостиТут — Администратор")
        self.resize(1100, 700)

        central = QWidget()
        h = QHBoxLayout()
        central.setLayout(h)
        self.setCentralWidget(central)

        # Левая панель меню
        sidebar = QWidget()
        sidebar.setFixedWidth(180)
        sidebar.setStyleSheet(f"background-color: {SIDEBAR_COLOR}; color: white;")
        sbv = QVBoxLayout()
        sbv.setContentsMargins(12, 12, 12, 12)
        title = QLabel("ГостиТут")
        title.setStyleSheet("color:white; font-weight:bold; font-size:16px;")
        sbv.addWidget(title)

        # Кнопки навигации
        btn_style = "text-align:center; padding:10px; border:none; color:#fff; background:transparent;"
        self.btn_main = QPushButton("Главная")
        self.btn_main.setStyleSheet(btn_style)
        self.btn_guests = QPushButton("Гости")
        self.btn_guests.setStyleSheet(btn_style)
        self.btn_rooms = QPushButton("Номера")
        self.btn_rooms.setStyleSheet(btn_style)
        self.btn_bookings = QPushButton("Брони")
        self.btn_bookings.setStyleSheet(btn_style)
        for b in (self.btn_main, self.btn_guests, self.btn_rooms, self.btn_bookings):
            b.setFixedHeight(36)
            sbv.addWidget(b)
        sbv.addStretch()

        # Информация о пользователе
        name_lbl = QLabel(f"{admin.get('first_name', '')}\n{admin.get('last_name', '')}")
        name_lbl.setAlignment(Qt.AlignmentFlag.AlignHCenter)
        name_lbl.setStyleSheet("color:#fff;")
        sbv.addWidget(name_lbl)
        logout = QPushButton("Выход")
        logout.setStyleSheet(
            "background:#3f3131; color:white; padding:6px; border-radius:6px;"
        )
        logout.clicked.connect(self.close)
        sbv.addWidget(logout)
        sidebar.setLayout(sbv)

        # Правая часть — вкладки
        self.stack = QStackedWidget()
        self.page_main = self.build_main_page()
        self.page_guests = self.build_guests_page()
        self.page_rooms = self.build_rooms_page()
        self.page_bookings = self.build_bookings_page()
        for p in (self.page_main, self.page_guests, self.page_rooms, self.page_bookings):
            self.stack.addWidget(p)

        # Навигация
        self.btn_main.clicked.connect(self.go_main)
        self.btn_guests.clicked.connect(
            lambda: self.stack.setCurrentWidget(self.page_guests)
        )
        self.btn_rooms.clicked.connect(
            lambda: self.stack.setCurrentWidget(self.page_rooms)
        )
        self.btn_bookings.clicked.connect(
            lambda: self.stack.setCurrentWidget(self.page_bookings)
        )

        h.addWidget(sidebar)
        h.addWidget(self.stack, 1)

    # -------- Главная --------

    def build_main_page(self):
        """Страница «Главная»."""
        w = QWidget()
        v = QVBoxLayout()
        v.setContentsMargins(18, 18, 18, 18)
        title = QLabel("Главная")
        title.setFont(TITLE_FONT)
        v.addWidget(title)

        # Кнопка смены статуса номера
        ctrl_h = QHBoxLayout()
        btn_status = QPushButton("Изменить статус")
        btn_status.clicked.connect(self.on_change_status)
        ctrl_h.addWidget(btn_status)
        ctrl_h.addStretch()
        v.addLayout(ctrl_h)

        # Легенда + картинка
        legend_col = QVBoxLayout()
        legend_col.setSpacing(8)
        legend_col.setAlignment(Qt.AlignmentFlag.AlignTop)
        for text, color in (
            ("Свободен", "#2ecc71"),
            ("Уборка", "#f1c40f"),
            ("Занят", "#e74c3c"),
            ("Бронь", "#95a5a6"),
        ):
            lbl = QLabel(f"● {text}")
            lbl.setStyleSheet(f"color: {color}; font-weight: 600;")
            lbl.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
            legend_col.addWidget(lbl)

        # Картинка, если есть
        if os.path.exists(MAIN_IMAGE_PATH):
            pix = QPixmap(MAIN_IMAGE_PATH).scaledToWidth(
                120, Qt.TransformationMode.SmoothTransformation
            )
            imw = QLabel()
            imw.setPixmap(pix)
            imw.setAlignment(Qt.AlignmentFlag.AlignTop)
            legend_col.addWidget(imw, alignment=Qt.AlignmentFlag.AlignTop)
        legend_col.addStretch()

        # Колонки с категориями номеров
        self.room_tiles = []
        cats = db.fetchall("SELECT id, name FROM room_types ORDER BY id")
        columns_h = QHBoxLayout()
        columns_h.setSpacing(18)
        for cat in cats:
            cat_id, cat_name = cat
            box = QVBoxLayout()
            hdr = QLabel(cat_name)
            hdr.setFont(SECTION_FONT)
            hdr.setAlignment(Qt.AlignmentFlag.AlignHCenter)
            box.addWidget(hdr)
            card = QVBoxLayout()
            rooms = db.fetchall(
                "SELECT number, status, id FROM rooms WHERE type_id=%s ORDER BY number",
                (cat_id,),
            )
            if not rooms:
                continue
            for r in rooms:
                number, status, rid = r
                tile = RoomTile(str(number) + "\n", status, rid)
                tile.clicked.connect(self.on_tile_clicked)
                self.room_tiles.append(tile)
                card.addWidget(tile)
            card.addStretch()
            container = QWidget()
            container.setLayout(card)
            container.setMinimumHeight(320)
            container.setStyleSheet(
                "background:white; border-radius:8px; padding:8px;"
            )
            box.addWidget(container)
            columns_h.addLayout(box)

        # Основная часть: слева номера, справа подсказка по цветам
        body_h = QHBoxLayout()
        body_h.addLayout(columns_h, 1)
        body_h.addSpacing(12)
        body_h.addLayout(legend_col)
        v.addLayout(body_h)
        w.setLayout(v)
        return w

    def go_main(self):
        """Обновляем главную вкладку, чтобы перечитать номера."""
        self.selected_tile = None
        if self.page_main:
            self.stack.removeWidget(self.page_main)
            self.page_main.deleteLater()
        self.page_main = self.build_main_page()
        self.stack.insertWidget(0, self.page_main)
        self.stack.setCurrentWidget(self.page_main)

    def on_tile_clicked(self, tile: RoomTile):
        """Выделяем выбранную карточку номера."""
        if self.selected_tile and self.selected_tile is not tile:
            try:
                self.selected_tile.set_selected(False)
            except RuntimeError:
                # старый виджет уже уничтожен при перестроении интерфейса
                pass
        self.selected_tile = tile
        tile.set_selected(True)

    def on_change_status(self):
        """Меняем статус выбранного номера."""
        if not self.selected_tile:
            QMessageBox.information(
                self, "Статус", "Сначала выберите номер (клик по карточке)."
            )
            return
        statuses = ["свободен", "уборка", "занят", "бронь"]
        current_index = (
            statuses.index(self.selected_tile.status)
            if self.selected_tile.status in statuses
            else 0
        )
        new_status, ok = QInputDialog.getItem(
            self,
            "Изменить статус",
            "Выберите новый статус:",
            statuses,
            current_index,
            False,
        )
        if not ok or not new_status:
            return
        try:
            db.execute(
                "UPDATE rooms SET status=%s WHERE id=%s",
                (new_status, self.selected_tile.room_id),
            )
        except Exception as e:
            QMessageBox.critical(self, "Ошибка БД", str(e))
            return
        self.selected_tile.set_status(new_status)
        QMessageBox.information(
            self, "Статус", f"Статус обновлен на «{new_status}»."
        )

    # -------- Гости --------

    def build_guests_page(self):
        """Страница «Гости»."""
        w = QWidget()
        v = QVBoxLayout()
        v.setContentsMargins(18, 18, 18, 18)
        title = QLabel("Гости")
        title.setFont(TITLE_FONT)
        v.addWidget(title)

        btn_h = QHBoxLayout()
        btn_add = QPushButton("Добавить гостя")
        btn_checkout = QPushButton("Выселить гостя")
        btn_report = QPushButton("Отчет по гостю")
        btn_h.addWidget(btn_add)
        btn_h.addWidget(btn_checkout)
        btn_h.addWidget(btn_report)
        btn_h.addStretch()
        v.addLayout(btn_h)

        self.guests_table = QTableWidget(0, 6)
        self.guests_table.setHorizontalHeaderLabels(
            ["ФИО", "Паспорт", "Номер", "Дата заезда", "Дата выезда", "Оплата"]
        )
        self.guests_table.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.Stretch
        )
        v.addWidget(self.guests_table)

        self.guests_table.itemDoubleClicked.connect(self.dialog_edit_guest)
        btn_add.clicked.connect(self.dialog_add_guest)
        btn_checkout.clicked.connect(self.action_checkout_guest)
        btn_report.clicked.connect(self.action_guest_report)

        self.reload_guests()
        w.setLayout(v)
        return w

    def reload_guests(self):
        self.guests_table.setRowCount(0)
        rows = db.fetchall(
            """
            SELECT g.id, g.first_name, g.last_name, g.passport_encrypted IS NOT NULL AS has_pass,
                   COALESCE(g.discount,0) AS discount,
                   b.room_id, b.date_from, b.date_to, b.total_price, b.status
            FROM guests g
            LEFT JOIN bookings b ON b.guest_id = g.id AND b.status IN ('active','completed')
            ORDER BY g.created_at DESC
            """
        )
        for r in rows:
            gid, fn, ln, has_pass, discount, room_id, dfrom, dto, price, bstatus = r
            row = self.guests_table.rowCount()
            self.guests_table.insertRow(row)
            fio_item = QTableWidgetItem(f"{fn} {ln}")
            fio_item.setData(Qt.ItemDataRole.UserRole, gid)
            self.guests_table.setItem(row, 0, fio_item)
            self.guests_table.setItem(
                row, 1, QTableWidgetItem("зашифровано" if has_pass else "")
            )
            self.guests_table.setItem(
                row, 2, QTableWidgetItem(str(room_id) if room_id else "")
            )
            self.guests_table.setItem(
                row, 3, QTableWidgetItem(str(dfrom) if dfrom else "")
            )
            self.guests_table.setItem(
                row, 4, QTableWidgetItem(str(dto) if dto else "")
            )
            pay_text = ""
            if price is not None:
                pay_text = str(price)
                if discount:
                    pay_text += f" (скидка {discount}%)"
            self.guests_table.setItem(row, 5, QTableWidgetItem(pay_text))

    def dialog_add_guest(self):
        """Окно добавления гостя и создания брони."""
        dlg = QDialog(self)
        dlg.setWindowTitle("Добавить гостя и создать бронь")
        form = QFormLayout()
        first = QLineEdit()
        last = QLineEdit()
        phone = QLineEdit()
        passport = QLineEdit()
        discount = QDoubleSpinBox()
        discount.setRange(0, 100)
        discount.setDecimals(2)
        discount.setValue(0.0)
        date_from = QDateEdit()
        date_from.setDate(QDate.currentDate())
        date_to = QDateEdit()
        date_to.setDate(QDate.currentDate().addDays(1))

        # Выбор только свободных номеров
        room_sel = QComboBox()
        rooms = db.fetchall(
            "SELECT id, number FROM rooms WHERE status='свободен' ORDER BY number"
        )
        for r in rooms:
            room_sel.addItem(r[1], r[0])

        form.addRow("Имя:", first)
        form.addRow("Фамилия:", last)
        form.addRow("Телефон:", phone)
        form.addRow("Паспорт (серия номер):", passport)
        form.addRow("Скидка, %:", discount)
        form.addRow("Номер:", room_sel)
        form.addRow("Заезд:", date_from)
        form.addRow("Выезд:", date_to)
        btn = QPushButton("Сохранить")

        def save():
            fn, ln = first.text().strip(), last.text().strip()
            pass_txt = passport.text().strip()
            if not fn or not ln or not pass_txt:
                QMessageBox.warning(dlg, "Ошибка", "Введите ФИО и паспорт")
                return

            # Шифруем паспорт
            nonce, ct = aes_encrypt(pass_txt.encode("utf-8"))
            try:
                with db.conn.cursor() as cur:
                    cur.execute(
                        """
                        INSERT INTO guests(first_name, last_name, phone, passport_encrypted, passport_iv, discount)
                        VALUES (%s,%s,%s,%s,%s,%s) RETURNING id
                        """,
                        (fn, ln, phone.text().strip(), ct, nonce, discount.value()),
                    )
                    gid = cur.fetchone()[0]

                    # Создаём бронь
                    room_id = room_sel.currentData()
                    dfrom = date_from.date().toPyDate()
                    dto = date_to.date().toPyDate()

                    # Проверяем, нет ли пересечений по датам
                    overlap = db.fetchone(
                        """
                        SELECT 1 FROM bookings
                        WHERE room_id=%s AND status='active' AND (%s < date_to) AND (%s > date_from)
                        """,
                        (room_id, dfrom, dto),
                    )
                    if overlap:
                        QMessageBox.warning(
                            dlg,
                            "Ошибка",
                            "Номер уже забронирован в указанный период",
                        )
                        db.conn.rollback()
                        return

                    # Считаем итоговую цену с учётом скидки
                    cur.execute(
                        """
                        SELECT rt.base_price
                        FROM rooms r JOIN room_types rt ON r.type_id=rt.id
                        WHERE r.id=%s
                        """,
                        (room_id,),
                    )
                    row_bp = cur.fetchone()
                    base_price = (
                        float(row_bp[0]) if row_bp and row_bp[0] is not None else 0.0
                    )
                    nights = (dto - dfrom).days
                    if nights < 1:
                        nights = 1
                    total = base_price * nights * (1 - discount.value() / 100.0)
                    cur.execute(
                        """
                        INSERT INTO bookings(room_id, guest_id, created_by, date_from, date_to, total_price)
                        VALUES (%s,%s,%s,%s,%s,%s) RETURNING id
                        """,
                        (room_id, gid, self.admin.get("id"), dfrom, dto, total),
                    )
                    bid = cur.fetchone()[0]

                    # Ставим номер в статус «бронь»
                    cur.execute("UPDATE rooms SET status='бронь' WHERE id=%s", (room_id,))
                    db.conn.commit()
                    QMessageBox.information(
                        dlg, "Готово", f"Гость добавлен, бронь id={bid}"
                    )
            except Exception as e:
                db.conn.rollback()
                QMessageBox.critical(self, "Ошибка БД", str(e))
            dlg.accept()
            self.reload_guests()
            self.reload_rooms()

        btn.clicked.connect(save)
        form.addRow(btn)
        dlg.setLayout(form)
        dlg.exec()

    def dialog_edit_guest(self):
        """Редактирование данных гостя и его активной брони."""
        row = self.guests_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Выбор", "Выберите гостя для редактирования")
            return
        fio_item = self.guests_table.item(row, 0)
        gid = fio_item.data(Qt.ItemDataRole.UserRole)
        if not gid:
            QMessageBox.warning(self, "Ошибка", "Не удалось определить гостя")
            return
        g = db.fetchone(
            """
            SELECT first_name, last_name, phone, email, passport_encrypted, passport_iv, COALESCE(discount,0)
            FROM guests WHERE id=%s
            """,
            (gid,),
        )
        if not g:
            QMessageBox.warning(self, "Ошибка", "Гость не найден")
            return
        fn_cur, ln_cur, phone_cur, email_cur, pen, piv, discount_cur = g
        passport_plain = ""
        try:
            if pen and piv:
                passport_plain = aes_decrypt(piv.tobytes(), pen.tobytes()).decode(
                    "utf-8"
                )
        except Exception:
            passport_plain = ""

        dlg = QDialog(self)
        dlg.setWindowTitle("Редактировать гостя")
        form = QFormLayout()
        first_name = QLineEdit(fn_cur or "")
        last_name = QLineEdit(ln_cur or "")
        phone = QLineEdit(phone_cur or "")
        email = QLineEdit(email_cur or "")
        passport = QLineEdit(passport_plain or "")
        discount = QDoubleSpinBox()
        discount.setRange(0, 100)
        discount.setDecimals(2)
        discount.setValue(float(discount_cur or 0))
        form.addRow("Имя", first_name)
        form.addRow("Фамилия", last_name)
        form.addRow("Телефон", phone)
        form.addRow("Email", email)
        form.addRow("Паспорт", passport)
        form.addRow("Скидка, %", discount)

        # Если есть активная бронь — даём выбрать другой номер
        active_booking = db.fetchone(
            """
            SELECT id, room_id, date_from, date_to
            FROM bookings
            WHERE guest_id=%s AND status='active'
            ORDER BY id DESC LIMIT 1
            """,
            (gid,),
        )
        room_combo = None
        old_room_id = None
        booking_dates = None
        if active_booking:
            bid, old_room_id, d_from, d_to = active_booking
            booking_dates = (d_from, d_to)
            room_combo = QComboBox()
            # Подбираем доступные номера с учётом дат
            rooms = db.fetchall(
                """
                SELECT r.id, r.number
                FROM rooms r
                WHERE r.id=%s
                   OR NOT EXISTS (
                        SELECT 1 FROM bookings b2
                        WHERE b2.room_id = r.id
                          AND b2.status='active'
                          AND b2.id<>%s
                          AND (%s < b2.date_to) AND (%s > b2.date_from)
                   )
                ORDER BY r.number
                """,
                (old_room_id, bid, d_from, d_to),
            )
            for rid, num in rooms:
                room_combo.addItem(num, rid)
                if rid == old_room_id:
                    room_combo.setCurrentIndex(room_combo.count() - 1)
            form.addRow("Номер (активная бронь)", room_combo)

        btn = QPushButton("Сохранить")

        def save():
            fn = first_name.text().strip()
            ln = last_name.text().strip()
            ph = phone.text().strip() or None
            em = email.text().strip() or None
            pass_txt = passport.text().strip() or None
            disc_val = discount.value()
            if not fn or not ln:
                QMessageBox.warning(self, "Ошибка", "Имя и фамилия обязательны")
                return
            new_room_id = room_combo.currentData() if room_combo else None
            try:
                pass_ct = None
                pass_iv = None
                if pass_txt:
                    pass_iv, pass_ct = aes_encrypt(pass_txt.encode("utf-8"))
                db.execute(
                    """
                    UPDATE guests
                    SET first_name=%s, last_name=%s, phone=%s, email=%s,
                        passport_encrypted=%s, passport_iv=%s, discount=%s
                    WHERE id=%s
                    """,
                    (fn, ln, ph, em, pass_ct, pass_iv, disc_val, gid),
                )

                # Если меняем номер или скидку — пересчитываем бронь и цену
                if active_booking and (
                    (new_room_id and new_room_id != old_room_id)
                    or disc_val != float(discount_cur or 0)
                ):
                    d_from, d_to = booking_dates
                    overlap = db.fetchone(
                        """
                        SELECT 1 FROM bookings
                        WHERE room_id=%s AND id<>%s AND status='active'
                          AND (%s < date_to) AND (%s > date_from)
                        """,
                        (new_room_id, bid, d_from, d_to),
                    )
                    if overlap:
                        QMessageBox.warning(
                            dlg, "Ошибка", "Номер занят/забронирован в эти даты"
                        )
                        db.conn.rollback()
                        return
                    room_for_price = new_room_id or old_room_id
                    with db.conn.cursor() as cur:
                        # Пересчитываем итоговую сумму
                        cur.execute(
                            """
                            SELECT rt.base_price
                            FROM rooms r JOIN room_types rt ON r.type_id=rt.id
                            WHERE r.id=%s
                            """,
                            (room_for_price,),
                        )
                        row_bp = cur.fetchone()
                        base_price = (
                            float(row_bp[0])
                            if row_bp and row_bp[0] is not None
                            else 0.0
                        )
                        nights = (d_to - d_from).days
                        if nights < 1:
                            nights = 1
                        total = base_price * nights * (1 - disc_val / 100.0)
                        if new_room_id and new_room_id != old_room_id:
                            cur.execute(
                                "UPDATE bookings SET room_id=%s, total_price=%s WHERE id=%s",
                                (new_room_id, total, bid),
                            )
                            # Обновляем статусы для старого и нового номера
                            cur.execute(
                                "UPDATE rooms SET status='свободен' WHERE id=%s",
                                (old_room_id,),
                            )
                            cur.execute(
                                "UPDATE rooms SET status='бронь' WHERE id=%s",
                                (new_room_id,),
                            )
                        else:
                            cur.execute(
                                "UPDATE bookings SET total_price=%s WHERE id=%s",
                                (total, bid),
                            )
                    db.conn.commit()
                QMessageBox.information(self, "Сохранено", "Данные гостя обновлены")
                dlg.accept()
                self.reload_guests()
                self.reload_rooms()
            except Exception as e:
                QMessageBox.critical(self, "Ошибка БД", str(e))

        btn.clicked.connect(save)
        form.addRow(btn)
        dlg.setLayout(form)
        dlg.exec()

    def action_checkout_guest(self):
        """Выселяем гостя и ставим номер в статус «уборка»."""
        row = self.guests_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Выбор", "Выберите гостя в таблице")
            return
        fio = self.guests_table.item(row, 0).text()
        parts = fio.split()
        if len(parts) < 2:
            QMessageBox.warning(self, "Ошибка", "Невозможно определить гостя")
            return
        first, last = parts[0], parts[1]
        g = db.fetchone(
            "SELECT id FROM guests WHERE first_name=%s AND last_name=%s ORDER BY id DESC LIMIT 1",
            (first, last),
        )
        if not g:
            QMessageBox.warning(self, "Ошибка", "Гость не найден")
            return
        gid = g[0]
        b = db.fetchone(
            "SELECT id, room_id FROM bookings WHERE guest_id=%s AND status='active' ORDER BY id DESC LIMIT 1",
            (gid,),
        )
        if not b:
            QMessageBox.information(self, "Инфо", "У гостя нет активной брони")
            return
        bid, room_id = b
        try:
            with db.conn.cursor() as cur:
                cur.execute(
                    "UPDATE bookings SET status='completed' WHERE id=%s", (bid,)
                )
                cur.execute(
                    "UPDATE rooms SET status='уборка' WHERE id=%s", (room_id,)
                )
                cur.execute(
                    """
                    INSERT INTO room_status_history(room_id, old_status, new_status, changed_by)
                    VALUES (%s,%s,%s,%s)
                    """,
                    (room_id, "занят", "уборка", self.admin.get("id")),
                )
                db.conn.commit()
            QMessageBox.information(
                self, "Готово", "Гость выселён, номер помечен как 'уборка'"
            )
        except Exception as e:
            db.conn.rollback()
            QMessageBox.critical(self, "Ошибка БД", str(e))
        self.reload_guests()
        self.reload_rooms()

    def action_guest_report(self):
        """Формируем docx‑отчёт по выбранному гостю."""
        row = self.guests_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Выбор", "Выберите гостя в таблице")
            return
        fio = self.guests_table.item(row, 0).text()
        parts = fio.split()
        first, last = parts[0], (parts[1] if len(parts) > 1 else "")

        g = db.fetchone(
            """
            SELECT id, first_name, last_name, phone, email, passport_encrypted, passport_iv, created_at
            FROM guests
            WHERE first_name=%s AND last_name=%s
            ORDER BY id DESC LIMIT 1
            """,
            (first, last),
        )
        if not g:
            QMessageBox.warning(self, "Ошибка", "Гость не найден")
            return
        gid, g_first, g_last, phone, email, pen, piv, created_at = g

        passport_plain = "не доступен"
        try:
            if pen and piv:
                passport_plain = aes_decrypt(piv.tobytes(), pen.tobytes()).decode(
                    "utf-8"
                )
        except Exception as e:
            passport_plain = f"Ошибка расшифровки: {e}"

        bookings = db.fetchall(
            """
            SELECT b.id, b.date_from, b.date_to, b.status, b.total_price, r.number
            FROM bookings b
            LEFT JOIN rooms r ON r.id=b.room_id
            WHERE b.guest_id=%s
            ORDER BY b.created_at DESC
            """,
            (gid,),
        )

        # Пытаемся создать docx; если нет библиотеки, предупредим
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            QMessageBox.warning(
                self,
                "Нет зависимости",
                "Установите пакет python-docx: pip install python-docx",
            )
            return

        doc = Document()
        doc.add_heading("Отчёт о проживании гостя", level=1)
        doc.add_paragraph(f"ФИО: {g_first} {g_last}")
        doc.add_paragraph(f"Паспорт: {passport_plain}")
        if phone:
            doc.add_paragraph(f"Телефон: {phone}")
        if email:
            doc.add_paragraph(f"E-mail: {email}")
        if created_at:
            doc.add_paragraph(f"Заведён в системе: {created_at}")

        doc.add_paragraph("")
        doc.add_heading("Бронирования", level=2)
        if not bookings:
            doc.add_paragraph("Нет данных о бронированиях.")
        else:
            table = doc.add_table(rows=1, cols=5)
            hdr = table.rows[0].cells
            hdr[0].text = "Номер"
            hdr[1].text = "Заезд"
            hdr[2].text = "Выезд"
            hdr[3].text = "Статус"
            hdr[4].text = "Сумма"
            for _, d_from, d_to, b_status, price, room_num in bookings:
                row_cells = table.add_row().cells
                row_cells[0].text = str(room_num or "")
                row_cells[1].text = str(d_from or "")
                row_cells[2].text = str(d_to or "")
                row_cells[3].text = b_status or ""
                row_cells[4].text = f"{price}" if price is not None else ""

        doc.add_paragraph("")
        footer = doc.add_paragraph('ООО "ГостиТут"')
        footer.runs[0].font.size = Pt(10)

        reports_dir = os.path.join(os.getcwd(), "reports")
        os.makedirs(reports_dir, exist_ok=True)
        filename = (
            f"guest_report_{gid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        filepath = os.path.join(reports_dir, filename)
        try:
            doc.save(filepath)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить отчёт: {e}")
            return

        QMessageBox.information(
            self, "Отчёт сформирован", f"Файл сохранён: {filepath}"
        )

    # -------- Номера --------

    def build_rooms_page(self):
        """Страница «Номера»."""
        w = QWidget()
        v = QVBoxLayout()
        v.setContentsMargins(18, 18, 18, 18)
        title = QLabel("Номера")
        title.setFont(TITLE_FONT)
        v.addWidget(title)

        btn_h = QHBoxLayout()
        btn_add_cat = QPushButton("Добавить категорию")
        btn_add_room = QPushButton("Добавить номер")
        btn_del_room = QPushButton("Удалить номер")
        btn_h.addWidget(btn_add_cat)
        btn_h.addWidget(btn_add_room)
        btn_h.addWidget(btn_del_room)
        btn_h.addStretch()
        v.addLayout(btn_h)

        # Таблица с номерами
        self.rooms_table = QTableWidget(0, 4)
        self.rooms_table.setHorizontalHeaderLabels(
            ["Номер", "Этаж", "Категория", "Статус"]
        )
        self.rooms_table.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.Stretch
        )
        v.addWidget(self.rooms_table)

        self.rooms_table.itemDoubleClicked.connect(self.dialog_edit_room)
        btn_add_cat.clicked.connect(self.dialog_add_category)
        btn_add_room.clicked.connect(self.dialog_add_room)
        btn_del_room.clicked.connect(self.action_delete_room)

        self.reload_rooms()
        w.setLayout(v)
        return w

    def reload_rooms(self):
        """Обновляем таблицу номеров и главную страницу, если она открыта."""
        self.selected_tile = None
        self.rooms_table.setRowCount(0)
        rows = db.fetchall(
            """
            SELECT r.id, r.number, r.floor, rt.name, r.status, rt.id AS type_id, rt.base_price
            FROM rooms r LEFT JOIN room_types rt ON r.type_id=rt.id
            ORDER BY r.number
            """
        )
        for r in rows:
            rid, number, floor, cat, status, type_id, base_price = r
            row = self.rooms_table.rowCount()
            self.rooms_table.insertRow(row)
            num_item = QTableWidgetItem(str(number))
            num_item.setData(
                Qt.ItemDataRole.UserRole,
                {
                    "id": rid,
                    "type_id": type_id,
                    "price": float(base_price) if base_price is not None else None,
                },
            )
            self.rooms_table.setItem(row, 0, num_item)
            self.rooms_table.setItem(
                row, 1, QTableWidgetItem(str(floor) if floor else "")
            )
            self.rooms_table.setItem(
                row, 2, QTableWidgetItem(cat if cat else "")
            )
            self.rooms_table.setItem(row, 3, QTableWidgetItem(status))

        # Если открыта «Главная» — обновляем плитки с номерами
        if self.stack.currentWidget() == self.page_main:
            self.stack.removeWidget(self.page_main)
            self.page_main = self.build_main_page()
            self.stack.insertWidget(0, self.page_main)
            self.stack.setCurrentWidget(self.page_main)

    def dialog_edit_room(self):
        """Редактирование одного номера."""
        row = self.rooms_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Выбор", "Выберите номер для редактирования")
            return
        num_item = self.rooms_table.item(row, 0)
        meta = num_item.data(Qt.ItemDataRole.UserRole) or {}
        rid = meta.get("id")
        type_id = meta.get("type_id")
        base_price = meta.get("price")
        if not rid:
            QMessageBox.warning(self, "Ошибка", "Не удалось определить номер")
            return
        number_cur = num_item.text()
        floor_cur = self.rooms_table.item(row, 1).text()

        dlg = QDialog(self)
        dlg.setWindowTitle(f"Редактировать номер {number_cur}")
        form = QFormLayout()
        number_edit = QLineEdit(number_cur)
        floor_spin = QSpinBox()
        floor_spin.setRange(-5, 100)
        try:
            floor_spin.setValue(int(floor_cur))
        except Exception:
            floor_spin.setValue(0)
        cat_combo = QComboBox()
        types = db.fetchall("SELECT id, name, base_price FROM room_types ORDER BY id")
        current_price = base_price
        for t_id, t_name, t_price in types:
            cat_combo.addItem(t_name, t_id)
            if t_id == type_id:
                current_price = t_price
        price_spin = QDoubleSpinBox()
        price_spin.setRange(0, 1_000_000)
        price_spin.setDecimals(2)
        if current_price is not None:
            price_spin.setValue(float(current_price))

        form.addRow("Категория", cat_combo)
        form.addRow("Номер", number_edit)
        form.addRow("Этаж", floor_spin)
        form.addRow("Цена категории", price_spin)
        btn = QPushButton("Сохранить")

        def save():
            num_new = number_edit.text().strip()
            if not num_new:
                QMessageBox.warning(dlg, "Ошибка", "Номер обязателен")
                return
            new_type_id = cat_combo.currentData()
            try:
                with db.conn.cursor() as cur:
                    cur.execute(
                        "UPDATE rooms SET number=%s, floor=%s, type_id=%s WHERE id=%s",
                        (num_new, floor_spin.value(), new_type_id, rid),
                    )
                    if new_type_id:
                        cur.execute(
                            "UPDATE room_types SET base_price=%s WHERE id=%s",
                            (price_spin.value(), new_type_id),
                        )
                db.conn.commit()
                QMessageBox.information(dlg, "Сохранено", "Номер обновлён")
                dlg.accept()
                self.reload_rooms()
                self.reload_guests()
            except Exception as e:
                db.conn.rollback()
                QMessageBox.critical(self, "Ошибка БД", str(e))

        btn.clicked.connect(save)
        form.addRow(btn)
        dlg.setLayout(form)
        dlg.exec()

    def dialog_add_category(self):
        """Добавление новой категории номера."""
        dlg = QDialog(self)
        dlg.setWindowTitle("Добавить категорию")
        form = QFormLayout()
        name = QLineEdit()
        desc = QLineEdit()
        price = QLineEdit()
        form.addRow("Название:", name)
        form.addRow("Описание:", desc)
        form.addRow("Базовая цена:", price)
        btn = QPushButton("Добавить")

        def save():
            n = name.text().strip()
            if not n:
                QMessageBox.warning(dlg, "Ошибка", "Введите название")
                return
            try:
                db.execute(
                    "INSERT INTO room_types(name, description, base_price) VALUES (%s,%s,%s)",
                    (n, desc.text().strip(), float(price.text() or 0)),
                )
                QMessageBox.information(dlg, "Готово", "Категория добавлена")
            except Exception as e:
                QMessageBox.critical(dlg, "Ошибка", str(e))
            dlg.accept()
            self.reload_rooms()

        btn.clicked.connect(save)
        form.addRow(btn)
        dlg.setLayout(form)
        dlg.exec()

    def dialog_add_room(self):
        """Добавление нового номера."""
        dlg = QDialog(self)
        dlg.setWindowTitle("Добавить номер")
        form = QFormLayout()
        number = QLineEdit()
        floor = QSpinBox()
        floor.setRange(0, 100)

        # Выбор категории номера
        cat = QComboBox()
        cats = db.fetchall("SELECT id, name FROM room_types ORDER BY id")
        for c in cats:
            cat.addItem(c[1], c[0])
        form.addRow("Номер:", number)
        form.addRow("Этаж:", floor)
        form.addRow("Категория:", cat)
        btn = QPushButton("Добавить")

        def save():
            if not number.text().strip():
                QMessageBox.warning(dlg, "Ошибка", "Введите номер")
                return
            try:
                db.execute(
                    "INSERT INTO rooms(number, floor, type_id) VALUES (%s,%s,%s)",
                    (number.text().strip(), floor.value(), cat.currentData()),
                )
                QMessageBox.information(dlg, "Готово", "Номер добавлен")
            except Exception as e:
                QMessageBox.critical(dlg, "Ошибка", str(e))
            dlg.accept()
            self.reload_rooms()

        btn.clicked.connect(save)
        form.addRow(btn)
        dlg.setLayout(form)
        dlg.exec()

    def action_delete_room(self):
        """Удаляем номер, при желании удаляем и категорию."""
        row = self.rooms_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Выбор", "Выберите строку")
            return
        num_item = self.rooms_table.item(row, 0)
        number = num_item.text()
        meta = num_item.data(Qt.ItemDataRole.UserRole) or {}
        room_id = meta.get("id")
        type_id = meta.get("type_id")

        # Окно подтверждения удаления с галочкой «удалить категорию»
        dlg = QDialog(self)
        dlg.setWindowTitle("Удалить номер")
        v = QVBoxLayout()
        v.addWidget(QLabel(f"Удалить номер {number}?"))
        chk = QCheckBox("Также удалить категорию этого номера из БД")
        v.addWidget(chk)
        btns = QHBoxLayout()
        btn_ok = QPushButton("Удалить")
        btn_cancel = QPushButton("Отмена")
        btns.addStretch()
        btns.addWidget(btn_cancel)
        btns.addWidget(btn_ok)
        v.addLayout(btns)
        dlg.setLayout(v)

        def do_delete():
            dlg.accept()
            if not room_id:
                QMessageBox.critical(self, "Ошибка", "Не удалось определить ID номера")
                return
            try:
                with db.conn.cursor() as cur:
                    cur.execute("DELETE FROM rooms WHERE id=%s", (room_id,))
                    if chk.isChecked() and type_id:
                        # проверим, остались ли ещё номера этой категории
                        cur.execute(
                            "SELECT COUNT(*) FROM rooms WHERE type_id=%s", (type_id,)
                        )
                        cnt = cur.fetchone()[0]
                        if cnt == 0:
                            cur.execute(
                                "DELETE FROM room_types WHERE id=%s", (type_id,)
                            )
                db.conn.commit()
                QMessageBox.information(self, "Готово", "Номер удалён")
                self.reload_rooms()
            except Exception as e:
                db.conn.rollback()
                QMessageBox.critical(self, "Ошибка", str(e))

        btn_ok.clicked.connect(do_delete)
        btn_cancel.clicked.connect(dlg.reject)
        dlg.exec()

    # -------- Брони --------

    def build_bookings_page(self):
        """Страница «Брони»."""
        w = QWidget()
        v = QVBoxLayout()
        v.setContentsMargins(18, 18, 18, 18)
        title = QLabel("Брони")
        title.setFont(TITLE_FONT)
        v.addWidget(title)
        btn_h = QHBoxLayout()
        btn_create = QPushButton("Создать бронь")
        btn_cancel = QPushButton("Отменить бронь")
        btn_h.addWidget(btn_create)
        btn_h.addWidget(btn_cancel)
        btn_h.addStretch()
        v.addLayout(btn_h)

        self.bookings_table = QTableWidget(0, 6)
        self.bookings_table.setHorizontalHeaderLabels(
            ["ID", "Номер", "Гость", "Заезд", "Выезд", "Статус"]
        )
        self.bookings_table.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.Stretch
        )
        v.addWidget(self.bookings_table)

        self.bookings_table.itemDoubleClicked.connect(self.dialog_edit_booking)
        btn_create.clicked.connect(self.dialog_create_booking)
        btn_cancel.clicked.connect(self.action_cancel_booking)

        self.reload_bookings()
        w.setLayout(v)
        return w

    def reload_bookings(self):
        self.bookings_table.setRowCount(0)
        rows = db.fetchall(
            """
            SELECT b.id, r.number, g.first_name||' '||g.last_name, b.date_from, b.date_to, b.status
            FROM bookings b
            LEFT JOIN rooms r ON r.id=b.room_id
            LEFT JOIN guests g ON g.id=b.guest_id
            ORDER BY b.date_from DESC
            """
        )
        for r in rows:
            row = self.bookings_table.rowCount()
            self.bookings_table.insertRow(row)
            for i, v in enumerate(r):
                self.bookings_table.setItem(row, i, QTableWidgetItem(str(v)))

    def dialog_edit_booking(self):
        """Редактирование выбранной брони."""
        row = self.bookings_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Выбор", "Выберите бронь для редактирования")
            return
        bid = int(self.bookings_table.item(row, 0).text())
        b = db.fetchone(
            """
            SELECT room_id, guest_id, date_from, date_to, status, total_price
            FROM bookings WHERE id=%s
            """,
            (bid,),
        )
        if not b:
            QMessageBox.warning(self, "Ошибка", "Бронь не найдена")
            return
        room_id, guest_id, d_from, d_to, status_cur, total_price = b

        dlg = QDialog(self)
        dlg.setWindowTitle(f"Редактировать бронь {bid}")
        form = QFormLayout()

        room_cb = QComboBox()
        rooms = db.fetchall("SELECT id, number, status FROM rooms ORDER BY number")
        for r in rooms:
            room_cb.addItem(f"{r[1]} ({r[2]})", r[0])
            if r[0] == room_id:
                room_cb.setCurrentIndex(room_cb.count() - 1)

        guest_cb = QComboBox()
        guests = db.fetchall(
            "SELECT id, first_name, last_name FROM guests ORDER BY id DESC"
        )
        for g in guests:
            guest_cb.addItem(f"{g[1]} {g[2]}", g[0])
            if g[0] == guest_id:
                guest_cb.setCurrentIndex(guest_cb.count() - 1)

        date_from = QDateEdit()
        date_to = QDateEdit()
        if isinstance(d_from, date):
            date_from.setDate(QDate(d_from.year, d_from.month, d_from.day))
        if isinstance(d_to, date):
            date_to.setDate(QDate(d_to.year, d_to.month, d_to.day))

        status_combo = QComboBox()
        statuses = ["active", "cancelled", "completed"]
        status_combo.addItems(statuses)
        if status_cur in statuses:
            status_combo.setCurrentText(status_cur)

        price_spin = QDoubleSpinBox()
        price_spin.setRange(0, 5_000_000)
        price_spin.setDecimals(2)
        if total_price is not None:
            price_spin.setValue(float(total_price))

        form.addRow("Номер", room_cb)
        form.addRow("Гость", guest_cb)
        form.addRow("Заезд", date_from)
        form.addRow("Выезд", date_to)
        form.addRow("Статус", status_combo)
        form.addRow("Сумма", price_spin)
        btn = QPushButton("Сохранить")

        def save():
            room_new = room_cb.currentData()
            guest_new = guest_cb.currentData()
            dfrom = date_from.date().toPyDate()
            dto = date_to.date().toPyDate()
            if dto < dfrom:
                QMessageBox.warning(dlg, "Ошибка", "Дата выезда раньше заезда")
                return

            # Проверяем пересечения только для активных броней
            if status_combo.currentText() == "active":
                overlap = db.fetchone(
                    """
                    SELECT 1 FROM bookings
                    WHERE room_id=%s AND id<>%s AND status='active'
                      AND (%s < date_to) AND (%s > date_from)
                    """,
                    (room_new, bid, dfrom, dto),
                )
                if overlap:
                    QMessageBox.warning(
                        dlg, "Ошибка", "Номер занят/забронирован в выбранные даты"
                    )
                    return

            try:
                with db.conn.cursor() as cur:
                    # пересчёт стоимости по текущей скидке гостя
                    cur.execute(
                        "SELECT COALESCE(discount,0) FROM guests WHERE id=%s",
                        (guest_new,),
                    )
                    disc_row = cur.fetchone()
                    disc_val = (
                        float(disc_row[0])
                        if disc_row and disc_row[0] is not None
                        else 0.0
                    )
                    cur.execute(
                        """
                        SELECT rt.base_price
                        FROM rooms r JOIN room_types rt ON r.type_id=rt.id
                        WHERE r.id=%s
                        """,
                        (room_new,),
                    )
                    row_bp = cur.fetchone()
                    base_price = (
                        float(row_bp[0]) if row_bp and row_bp[0] is not None else 0.0
                    )
                    nights = (dto - dfrom).days
                    if nights < 1:
                        nights = 1
                    total = base_price * nights * (1 - disc_val / 100.0)
                    cur.execute(
                        """
                        UPDATE bookings
                        SET room_id=%s, guest_id=%s, date_from=%s, date_to=%s,
                            status=%s, total_price=%s
                        WHERE id=%s
                        """,
                        (
                            room_new,
                            guest_new,
                            dfrom,
                            dto,
                            status_combo.currentText(),
                            total,
                            bid,
                        ),
                    )

                    # Обновим статус комнаты в зависимости от статуса брони
                    if status_combo.currentText() == "active":
                        cur.execute(
                            "UPDATE rooms SET status='бронь' WHERE id=%s", (room_new,)
                        )
                    elif status_combo.currentText() in ("cancelled", "completed"):
                        cur.execute(
                            "UPDATE rooms SET status='свободен' WHERE id=%s",
                            (room_new,),
                        )
                db.conn.commit()
                QMessageBox.information(dlg, "Сохранено", "Бронь обновлена")
                dlg.accept()
                self.reload_bookings()
                self.reload_guests()
                self.reload_rooms()
            except Exception as e:
                db.conn.rollback()
                QMessageBox.critical(self, "Ошибка БД", str(e))

        btn.clicked.connect(save)
        form.addRow(btn)
        dlg.setLayout(form)
        dlg.exec()

    def dialog_create_booking(self):
        """Создание новой брони."""
        dlg = QDialog(self)
        dlg.setWindowTitle("Создать бронь")
        form = QFormLayout()

        # Выбор номера, который сейчас свободен
        room_cb = QComboBox()
        rooms = db.fetchall(
            """
            SELECT id, number, status
            FROM rooms
            WHERE status NOT IN ('занят','бронь')
            ORDER BY number
            """
        )
        for r in rooms:
            room_cb.addItem(f"{r[1]} ({r[2]})", r[0])

        # Выбор гостя из уже заведённых
        guest_cb = QComboBox()
        guests = db.fetchall(
            "SELECT id, first_name, last_name FROM guests ORDER BY id DESC"
        )
        for g in guests:
            guest_cb.addItem(f"{g[1]} {g[2]}", g[0])

        date_from = QDateEdit()
        date_from.setDate(QDate.currentDate())
        date_to = QDateEdit()
        date_to.setDate(QDate.currentDate().addDays(1))
        form.addRow("Номер:", room_cb)
        form.addRow("Гость:", guest_cb)
        form.addRow("Заезд:", date_from)
        form.addRow("Выезд:", date_to)
        btn = QPushButton("Создать")

        def create():
            room_id = room_cb.currentData()
            guest_id = guest_cb.currentData()
            dfrom = date_from.date().toPyDate()
            dto = date_to.date().toPyDate()
            if dfrom >= dto:
                QMessageBox.warning(
                    dlg, "Ошибка", "Дата выезда должна быть позже даты заезда"
                )
                return

            # Проверяем, нет ли пересечений по датам
            overlap = db.fetchone(
                """
                SELECT 1 FROM bookings
                WHERE room_id=%s AND status='active' AND (%s < date_to) AND (%s > date_from)
                """,
                (room_id, dfrom, dto),
            )
            if overlap:
                QMessageBox.warning(
                    dlg, "Ошибка", "Номер занят/забронирован в выбранные даты"
                )
                return

            try:
                with db.conn.cursor() as cur:
                    cur.execute(
                        """
                        INSERT INTO bookings(room_id, guest_id, created_by, date_from, date_to, total_price)
                        VALUES (%s,%s,%s,%s,%s,%s) RETURNING id
                        """,
                        (room_id, guest_id, self.admin.get("id"), dfrom, dto, 0.0),
                    )
                    bid = cur.fetchone()[0]
                    cur.execute(
                        "UPDATE rooms SET status='бронь' WHERE id=%s", (room_id,)
                    )
                    db.conn.commit()
                QMessageBox.information(dlg, "Готово", f"Бронь создана id={bid}")
            except Exception as e:
                db.conn.rollback()
                QMessageBox.critical(self, "Ошибка БД", str(e))
            dlg.accept()
            self.reload_bookings()
            self.reload_rooms()

        btn.clicked.connect(create)
        form.addRow(btn)
        dlg.setLayout(form)
        dlg.exec()

    def action_cancel_booking(self):
        """Отмена выбранной брони и освобождение номера."""
        row = self.bookings_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Выбор", "Выберите бронь")
            return
        bid = int(self.bookings_table.item(row, 0).text())
        try:
            with db.conn.cursor() as cur:
                # получить room_id
                cur.execute("SELECT room_id FROM bookings WHERE id=%s", (bid,))
                row = cur.fetchone()
                room_id = row[0] if row else None
                cur.execute(
                    "UPDATE bookings SET status='cancelled' WHERE id=%s", (bid,)
                )
                if room_id:
                    cur.execute(
                        "UPDATE rooms SET status='свободен' WHERE id=%s", (room_id,)
                    )
                db.conn.commit()
            QMessageBox.information(self, "Готово", "Бронь отменена")
        except Exception as e:
            db.conn.rollback()
            QMessageBox.critical(self, "Ошибка БД", str(e))
        self.reload_bookings()
        self.reload_rooms()


